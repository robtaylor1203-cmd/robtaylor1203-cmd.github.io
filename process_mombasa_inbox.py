import json
import datetime
import zipfile
import shutil
import requests
import re
from pathlib import Path
from pipeline_utils import generate_manifest

# Enhanced imports with fallback options
try:
    import fitz  # PyMuPDF for PDFs
    PYMUPDF_AVAILABLE = True
    print("PyMuPDF loaded successfully")
except ImportError as e:
    PYMUPDF_AVAILABLE = False
    print(f"Warning: PyMuPDF not available: {e}")

try:
    import docx  # python-docx for .docx files
    DOCX_AVAILABLE = True
    print("python-docx loaded successfully")
except ImportError as e:
    DOCX_AVAILABLE = False
    print(f"Warning: python-docx not available: {e}")

try:
    import openpyxl  # openpyxl for .xlsx files
    OPENPYXL_AVAILABLE = True
    print("openpyxl loaded successfully")
except ImportError as e:
    OPENPYXL_AVAILABLE = False
    print(f"Warning: openpyxl not available: {e}")

try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_AVAILABLE = True
    print("Playwright loaded successfully")
except ImportError as e:
    PLAYWRIGHT_AVAILABLE = False
    print(f"Warning: Playwright not available: {e}")

from urllib.parse import urljoin, urlparse

# --- Configuration ---
REPO_ROOT = Path(__file__).resolve().parent
LOCATION = "mombasa"
LANDING_PAGE_URL = "https://www.tbeal.net/tbea-market-report/"
BASE_URL = "https://www.tbeal.net"

# Standardized directories
INBOX_PATH = REPO_ROOT / "Inbox" / LOCATION
INBOX_PATH.mkdir(parents=True, exist_ok=True)

OUTPUT_DIR = REPO_ROOT / "source_reports" / LOCATION
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

ARCHIVE_PATH = INBOX_PATH / "archive"
ARCHIVE_PATH.mkdir(parents=True, exist_ok=True)

TEMP_DOWNLOAD_DIR = REPO_ROOT / "temp_downloads" / LOCATION
TEMP_DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Enhanced timeouts
MAX_TIMEOUT = 3600000  # 60 minutes
DISCOVERY_TIMEOUT = 600000  # 10 minutes
USER_AGENT = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'

# Reference data for calculating the sale number
REFERENCE_DATE = datetime.date(2025, 9, 1)
REFERENCE_SALE_NO = 34

def calculate_upcoming_sale_no():
    """Calculates the sale number for the upcoming Monday."""
    today = datetime.date.today()
    start_of_week = today - datetime.timedelta(days=today.weekday())
    weeks_passed = (start_of_week - REFERENCE_DATE).days // 7
    
    if weeks_passed < 0:
        print("Warning: Reference date is in the future or calculation error.")
        return REFERENCE_SALE_NO 

    current_sale_no = REFERENCE_SALE_NO + weeks_passed
    print(f"Calculated Current/Upcoming Sale Number: {current_sale_no}")
    return current_sale_no

def extract_text_from_file(file_path: Path):
    """Intelligently extracts text from a file based on its extension."""
    extension = file_path.suffix.lower()
    text = ""

    print(f"  - Reading file: {file_path.name}")

    try:
        if extension == ".txt":
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
                
        elif extension == ".pdf":
            if not PYMUPDF_AVAILABLE:
                return f"PyMuPDF not available - cannot process PDF: {file_path.name}"
            
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
                    
        elif extension == ".docx":
            if not DOCX_AVAILABLE:
                return f"python-docx not available - cannot process DOCX: {file_path.name}"
            
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'
                
        elif extension == ".doc":
            if not DOCX_AVAILABLE:
                return f"python-docx not available - cannot process DOC: {file_path.name}"
            
            # Try to read .doc as .docx (sometimes works)
            try:
                doc = docx.Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + '\n'
            except:
                text = f"Could not process .doc file: {file_path.name}. Consider converting to .docx"
                
        elif extension in [".xlsx", ".xls"]:
            if not OPENPYXL_AVAILABLE:
                return f"openpyxl not available - cannot process Excel: {file_path.name}"
            
            workbook = openpyxl.load_workbook(file_path)
            for sheetname in workbook.sheetnames:
                sheet = workbook[sheetname]
                text += f"--- Sheet: {sheetname} ---\n"
                for row in sheet.iter_rows(values_only=True):
                    text += "\t".join(str(cell) for cell in row if cell is not None) + "\n"
        else:
            text = f"Unsupported file type: {extension}"
            print(f"    Warning: {text}")
            
    except Exception as e:
        print(f"    Error extracting text from {file_path.name}: {e}")
        text = f"Error processing file: {e}"
        
    return text

def extract_sale_and_year_from_title(title):
    """Extract sale number and year from TBEAL report titles."""
    sale_match = re.search(r'Sale\s+(\d{1,2})', title, re.IGNORECASE)
    year_match = re.search(r'(\d{4})', title)
    
    sale_number = None
    year = None
    
    if sale_match:
        sale_number = sale_match.group(1).zfill(2)
    
    if year_match:
        year = year_match.group(1)
    
    # Fallback: current year if no year found
    if not year:
        year = str(datetime.datetime.now().year)
        
    return sale_number, year

def get_all_tbeal_reports(page):
    """Discover all available TBEAL market reports with enhanced waiting."""
    print("Discovering all available TBEAL market reports...")
    
    page.wait_for_load_state("networkidle", timeout=DISCOVERY_TIMEOUT)
    page.wait_for_timeout(15000)
    
    all_reports = []
    
    try:
        # Enhanced strategy: Look for various patterns of year links
        year_patterns = [
            'a[href*="market-report"]',
            'a[href*="tbea-market-report"]', 
            'a[href*="2025"]',
            'a[href*="2024"]',
            'a[href*="2023"]'
        ]
        
        for pattern in year_patterns:
            try:
                year_links = page.locator(pattern).all()
                print(f"Found {len(year_links)} links with pattern: {pattern}")
                
                for link in year_links:
                    try:
                        href = link.get_attribute('href')
                        text = link.inner_text().strip()
                        
                        if href and text:
                            # Check if it's a year category or direct download
                            if any(year in text for year in ['2025', '2024', '2023', '2022', '2021']):
                                year_match = re.search(r'(20\d{2})', text)
                                if year_match and 'download' not in href.lower():
                                    all_reports.append({
                                        'type': 'year_category',
                                        'year': year_match.group(1),
                                        'href': href,
                                        'text': text
                                    })
                                    print(f"  Found year category: {text}")
                            
                            # Check for direct download links
                            if 'download' in href.lower() and ('market' in text.lower() or 'sale' in text.lower()):
                                sale_number, year = extract_sale_and_year_from_title(text)
                                if sale_number:
                                    all_reports.append({
                                        'type': 'direct_download',
                                        'sale_number': sale_number,
                                        'year': year,
                                        'href': href,
                                        'title': text
                                    })
                                    print(f"  Found direct download: {text}")
                    except Exception as e:
                        continue
            except Exception as e:
                print(f"Pattern {pattern} failed: {e}")
                continue
    
    except Exception as e:
        print(f"Error discovering reports: {e}")
    
    # Direct access strategy for recent years
    print("\nTrying direct access strategy for recent years...")
    direct_access_years = [
        {"year": "2025", "url": "https://www.tbeal.net/download-category/tbea-market-report-2025/"},
        {"year": "2024", "url": "https://www.tbeal.net/download-category/tbea-market-report-2024/"},
        {"year": "2023", "url": "https://www.tbeal.net/download-category/tbea-market-report-2023/"}
    ]
    
    found_years = {item['year'] for item in all_reports if item.get('type') == 'year_category'}
    
    for year_info in direct_access_years:
        year = year_info['year']
        url = year_info['url']
        
        if year not in found_years:
            print(f"Attempting direct access to {year} via: {url}")
            all_reports.append({
                'type': 'year_category',
                'year': year,
                'href': url,
                'text': f'TBEA Market Report {year} (Direct Access)'
            })
            found_years.add(year)
    
    print(f"Discovered {len(all_reports)} potential reports/categories")
    return all_reports

def explore_year_category(context, category_info):
    """Explore a specific year category to find individual reports."""
    year = category_info['year']
    year_url = category_info['href']
    
    if not year_url.startswith('http'):
        year_url = urljoin(BASE_URL, year_url)
    
    print(f"Exploring {year} category: {year_url}")
    
    page = None
    try:
        page = context.new_page()
        page.set_default_timeout(MAX_TIMEOUT)
        page.add_init_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
        
        page.goto(year_url, wait_until="networkidle", timeout=300000)
        page.wait_for_timeout(8000)
        
        download_links = page.locator('a[href*="download"]').all()
        
        year_reports = []
        for link in download_links:
            try:
                href = link.get_attribute('href')
                text = link.inner_text().strip()
                
                if href and ('market-report' in href.lower() or 'sale' in text.lower()):
                    sale_number, report_year = extract_sale_and_year_from_title(text)
                    if sale_number:
                        download_url = urljoin(BASE_URL, href) if not href.startswith('http') else href
                        year_reports.append({
                            'sale_number': sale_number,
                            'year': report_year or year,
                            'href': href,
                            'title': text,
                            'download_url': download_url
                        })
                        print(f"  Found report: Sale {sale_number} ({report_year or year})")
            except Exception as e:
                continue
        
        print(f"Found {len(year_reports)} reports in {year}")
        return year_reports
        
    except Exception as e:
        print(f"Error exploring {year} category: {e}")
        return []
    finally:
        if page and not page.is_closed():
            page.close()

def download_word_document(url, filename):
    """Download a Word document from TBEAL."""
    print(f"Downloading Word document: {filename}")
    
    try:
        headers = {
            'User-Agent': USER_AGENT,
            'Accept': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document,application/msword,*/*'
        }
        
        response = requests.get(url, headers=headers, timeout=120)
        response.raise_for_status()
        
        # Determine file extension from content type or URL
        content_type = response.headers.get('content-type', '').lower()
        if 'wordprocessingml' in content_type:
            extension = '.docx'
        elif 'msword' in content_type:
            extension = '.doc'
        else:
            # Fallback: try to determine from URL
            if url.lower().endswith('.docx'):
                extension = '.docx'
            elif url.lower().endswith('.doc'):
                extension = '.doc'
            else:
                extension = '.docx'  # Default assumption
        
        # Ensure filename has correct extension
        if not filename.lower().endswith(('.doc', '.docx')):
            filename = filename + extension
        
        file_path = TEMP_DOWNLOAD_DIR / filename
        
        with open(file_path, 'wb') as f:
            f.write(response.content)
        
        print(f"Successfully downloaded to: {file_path}")
        return file_path
        
    except Exception as e:
        print(f"Error downloading {filename}: {e}")
        return None

def scrape_tbeal_all_reports():
    """Main function to scrape all available TBEAL reports."""
    print(f"Starting TBEAL Market Reports scraper - ALL REPORTS (Enhanced Version)")
    print("\nIMPORTANT: This will discover and process ALL available TBEAL reports.\n")

    if not PLAYWRIGHT_AVAILABLE:
        print("ERROR: Playwright is not available. Please install:")
        print("  pip install playwright")
        print("  playwright install chromium")
        return

    with sync_playwright() as p:
        browser = p.chromium.launch(
            headless=True, 
            timeout=MAX_TIMEOUT,
            args=['--no-sandbox', '--disable-setuid-sandbox', '--disable-dev-shm-usage']
        )
        context = browser.new_context(
            user_agent=USER_AGENT,
            viewport={'width': 1920, 'height': 1080}
        )
        
        try:
            page = context.new_page()
            page.set_default_timeout(DISCOVERY_TIMEOUT)
            page.add_init_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
            
            print("Navigating to TBEAL market reports page...")
            page.goto(LANDING_PAGE_URL, wait_until="networkidle", timeout=DISCOVERY_TIMEOUT)
            
            initial_reports = get_all_tbeal_reports(page)
            page.close()
            
            if not initial_reports:
                print("Warning: Could not find any TBEAL market reports.")
                return
            
            # Explore year categories to get individual reports
            all_individual_reports = []
            
            for item in initial_reports:
                if item['type'] == 'year_category':
                    year_reports = explore_year_category(context, item)
                    all_individual_reports.extend(year_reports)
                elif item['type'] == 'direct_download':
                    all_individual_reports.append(item)
            
            if not all_individual_reports:
                print("Warning: No individual reports found.")
                return
            
            # Sort by year and sale number (newest first)
            all_individual_reports.sort(key=lambda x: (int(x['year']), int(x['sale_number'])), reverse=True)
            
            # Show summary
            years_found = {}
            for report in all_individual_reports:
                year = report['year']
                years_found[year] = years_found.get(year, 0) + 1
            
            print(f"\nTBEAL Discovery Summary:")
            for year in sorted(years_found.keys(), reverse=True):
                print(f"  {year}: {years_found[year]} reports")
            
            # Limit to reasonable number for processing
            max_reports = 20  # Process up to 20 latest reports
            reports_to_process = all_individual_reports[:max_reports]
            
            print(f"\nProcessing {len(reports_to_process)} latest reports:")
            for report in reports_to_process:
                print(f"  - Sale {report['sale_number']} ({report['year']}): {report['title']}")

            # Process each report
            for i, report in enumerate(reports_to_process):
                sale_number = report['sale_number']
                year = report['year']
                title = report['title']
                download_url = report.get('download_url', report['href'])
                
                if not download_url.startswith('http'):
                    download_url = urljoin(BASE_URL, download_url)
                
                print(f"\n{'='*20} Processing Report {i+1}/{len(reports_to_process)}: Sale {sale_number} ({year}) {'='*20}")
                process_single_tbeal_report(sale_number, year, title, download_url)

        except Exception as e:
            print(f"Critical error occurred: {e}")
        finally:
            browser.close()
            print("Browser closed.")

def process_single_tbeal_report(sale_number, year, title, download_url):
    """Process a single TBEAL market report."""
    
    try:
        safe_filename = f"tbeal_market_report_S{sale_number}_{year}"
        
        downloaded_file = download_word_document(download_url, safe_filename)
        
        if not downloaded_file or not downloaded_file.exists():
            print(f"Failed to download report for Sale {sale_number} ({year})")
            return
        
        print(f"Extracting text from downloaded Word document...")
        extracted_text = extract_text_from_file(downloaded_file)
        
        if not extracted_text or extracted_text.startswith("Error"):
            print(f"Failed to extract text from Sale {sale_number} ({year})")
            return
        
        output_data = {
            "report_title": title,
            "sale_number": sale_number,
            "year": year,
            "source_url": download_url,
            "retrieval_date": datetime.datetime.now(datetime.timezone.utc).isoformat(),
            "file_type": "Word Document",
            "content_length": len(extracted_text),
            "extracted_text": extracted_text
        }
        
        print(f"Extracted {len(extracted_text)} characters from Sale {sale_number} ({year})")
        
        # Clean up downloaded file
        downloaded_file.unlink()
        print(f"Cleaned up temporary download")
        
    except Exception as e:
        print(f"\nERROR: Failed during TBEAL processing for Sale {sale_number} ({year}). Details: {e}")
        return

    if output_data and extracted_text.strip():
        sale_suffix = str(sale_number).zfill(2)
        filename = f"tbeal_market_report_enhanced_W{sale_suffix}_{year}.json"
        output_path = OUTPUT_DIR / filename

        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(output_data, f, indent=4, ensure_ascii=False, default=str)
            
            print(f"SUCCESS! TBEAL report processed for Sale {sale_number} ({year}).")
            print(f"Successfully saved report data to {output_path}")
            
            generate_manifest(REPO_ROOT, LOCATION, f"W{sale_suffix}_{year}", currency="KES")

        except Exception as e:
            print(f"Error saving file or generating manifest: {e}")
    else:
        print(f"TBEAL processing finished for Sale {sale_number} ({year}) without valid data.")

def process_local_sale_files():
    """Find and process local files for a given sale."""
    print(f"Starting Local File Processor for Mombasa Sales (Enhanced Version)")

    try:
        sale_no = calculate_upcoming_sale_no()
        
        # Define file paths based on calculated sale number
        zip_filename = f"{sale_no}.zip"
        txt_filename = f"Mombasa Sale {sale_no}.txt"
        current_year = datetime.date.today().year
        pdf_filename = f"Final Market Report Sale {sale_no} {current_year}.pdf"

        zip_filepath = INBOX_PATH / zip_filename
        txt_filepath = INBOX_PATH / txt_filename
        pdf_filepath = INBOX_PATH / pdf_filename
        
        temp_extract_path = INBOX_PATH / f"temp_{sale_no}"

        final_data = {
            "sale_number": sale_no,
            "retrieval_date": datetime.datetime.now(datetime.timezone.utc).isoformat(),
            "main_report_text": None,
            "final_market_report_pdf_text": None,
            "zip_attachments_content": {}
        }
        
        files_found = False

        # Process the Main Text File
        print(f"\n[1/3] Processing main text file: {txt_filename}")
        if txt_filepath.exists():
            final_data["main_report_text"] = extract_text_from_file(txt_filepath)
            print("  -> Success.")
            shutil.move(str(txt_filepath), ARCHIVE_PATH / txt_filepath.name)
            files_found = True
        else:
            print(f"  -> Warning: File not found.")
            final_data["main_report_text"] = "File not found."

        # Process the Final Market Report PDF
        print(f"\n[2/3] Processing Final Market Report PDF: {pdf_filename}")
        if pdf_filepath.exists():
            final_data["final_market_report_pdf_text"] = extract_text_from_file(pdf_filepath)
            print("  -> Success.")
            shutil.move(str(pdf_filepath), ARCHIVE_PATH / pdf_filepath.name)
            files_found = True
        else:
            print(f"  -> Warning: File not found.")
            final_data["final_market_report_pdf_text"] = "File not found."

        # Process the ZIP File
        print(f"\n[3/3] Processing ZIP file: {zip_filename}")
        if zip_filepath.exists():
            temp_extract_path.mkdir(parents=True, exist_ok=True)
            print(f"  - Extracting to temporary directory...")
            
            try:
                with zipfile.ZipFile(zip_filepath, 'r') as zip_ref:
                    zip_ref.extractall(temp_extract_path)
                
                print(f"  - Reading contents of extracted files...")
                for filepath in temp_extract_path.iterdir():
                    if filepath.is_file():
                        content = extract_text_from_file(filepath)
                        final_data["zip_attachments_content"][filepath.name] = content
                
                print("  - Cleaning up temporary files...")
                shutil.rmtree(temp_extract_path)
                print("  -> Success.")
                shutil.move(str(zip_filepath), ARCHIVE_PATH / zip_filepath.name)
                files_found = True
                
            except Exception as e:
                print(f"  -> Error processing ZIP file: {e}")
                final_data["zip_attachments_content"] = {"error": str(e)}
                
        else:
            print(f"  -> Warning: ZIP file not found.")
            final_data["zip_attachments_content"] = {"error": "File not found."}

        if files_found:
            file_prefix = "mombasa_processed_enhanced"
            sale_suffix = str(sale_no).zfill(2)
            output_filename = f"{file_prefix}_W{sale_suffix}.json"
            output_path = OUTPUT_DIR / output_filename
            
            print(f"\nSaving all extracted data to: {output_path}")
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(final_data, f, indent=4, ensure_ascii=False, default=str)

            generate_manifest(REPO_ROOT, LOCATION, sale_suffix, currency="KES")
            print("\n--- LOCAL FILE PROCESSING COMPLETE ---")
        else:
            print("\n--- PROCESS COMPLETE: No files found matching the calculated sale number. ---")

    except Exception as e:
        print(f"ERROR: An unexpected error occurred: {e}")

def check_dependencies():
    """Check if all required dependencies are available."""
    print("Checking dependencies...")
    
    missing = []
    available = []
    
    if PYMUPDF_AVAILABLE:
        available.append("PyMuPDF (PDF processing)")
    else:
        missing.append("PyMuPDF")
    
    if DOCX_AVAILABLE:
        available.append("python-docx (Word processing)")
    else:
        missing.append("python-docx")
    
    if OPENPYXL_AVAILABLE:
        available.append("openpyxl (Excel processing)")
    else:
        missing.append("openpyxl")
    
    if PLAYWRIGHT_AVAILABLE:
        available.append("Playwright (web automation)")
    else:
        missing.append("playwright")
    
    print("Available:")
    for dep in available:
        print(f"  ✓ {dep}")
    
    if missing:
        print(f"\nMissing (optional):")
        for dep in missing:
            print(f"  - {dep}")
        print(f"\nTo install missing dependencies:")
        print("  pip install PyMuPDF python-docx openpyxl playwright requests")
        print("  playwright install chromium")
    
    return True  # Continue even with missing optional dependencies

if __name__ == "__main__":
    print("=== COMPREHENSIVE MOMBASA MARKET REPORTS PROCESSOR (ENHANCED) ===")
    print("This script processes:")
    print("1. TBEAL web reports (Word/HTML documents)")
    print("2. Local inbox files (ZIP, PDF, TXT)")
    print("")
    
    check_dependencies()
    
    # Run the comprehensive TBEAL scraper for ALL reports
    print("\nPHASE 1: TBEAL Web Reports")
    try:
        scrape_tbeal_all_reports()
    except Exception as e:
        print(f"Error in TBEAL processing: {e}")
    
    print("\n" + "="*60)
    
    # Run local file processing
    print("PHASE 2: Local File Processing")
    try:
        process_local_sale_files()
    except Exception as e:
        print(f"Error in local file processing: {e}")
    
    print("\n=== ALL MOMBASA PROCESSING COMPLETE ===")
    
    # Uncomment individual phases if you want to run them separately:
    # scrape_tbeal_all_reports()        # Only TBEAL web scraping
    # process_local_sale_files()        # Only local file processing
